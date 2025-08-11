from docx import Document
import re
from typing import Dict

class DocxTemplateProcessor:
    def __init__(self, template_path: str):
        self.template_path = template_path
        self.variables = {}
        self.pattern = re.compile(r"\$[a-zA-Z0-9_\u4e00-\u9fa5]+\$")  # 匹配 $变量名$，支持中文
        self.doc = Document(template_path)

    def extract_variables(self) -> Dict[str, str]:
        """从文档中提取所有 $变量$ 并去重"""
        found = set()
        for para in self.doc.paragraphs:
            found.update(self.pattern.findall(para.text))
        self.variables = {var: "" for var in found}
        return self.variables

    def set_variable(self, var_name: str, value: str):
        """设置变量值"""
        var = f"${var_name}$"
        if var in self.variables:
            self.variables[var] = value
        else:
            print(f"[WARN] 变量 {var} 不存在于模板中。")

    def apply_variables(self, output_path: str):
        """将变量替换写入新的 Word 文档"""
        new_doc = Document(self.template_path)
        for para in new_doc.paragraphs:
            for var, val in self.variables.items():
                if var in para.text:
                    para.text = para.text.replace(var, val)

        new_doc.save(output_path)
        print(f"[INFO] 替换完成，输出文件保存到：{output_path}")


# 使用示例
if __name__ == "__main__":
    template_path = "模板.docx"
    output_path = "替换后文档.docx"

    processor = DocxTemplateProcessor(template_path)
    vars_found = processor.extract_variables()

    print("在文档中发现的变量：")
    for var in vars_found:
        print(f"  {var}")

    # 设置变量值（你可以根据实际接口接收来自外部输入）
    processor.set_variable("缺陷统计", "总数：25项")
    processor.set_variable("检查日期", "2025-07-18")
    processor.set_variable("图片数量", "共10张")

    # 替换并保存新文档
    processor.apply_variables(output_path)
